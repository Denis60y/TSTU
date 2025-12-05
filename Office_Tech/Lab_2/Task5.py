import docx 
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'

table = doc.add_table(rows=1, cols=1)
table.style = 'TableGrid'
table.autofit = False  

row = table.rows[0]
row.height = Cm(6) 

cell = table.cell(0, 0)
cell.width = Cm(9)

for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

text = cell.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(3.6)
add = text.add_run("___________________")
add.font.size = Pt(14)

text = cell.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(5.5)
add = text.add_run("(учредитель)")
add.font.size = Pt(6)

text = cell.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(3.6)
add = text.add_run("_____________________________")
add.font.size = Pt(9)

text = cell.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(3.6)
add = text.add_run("___________________")
add.font.size = Pt(14)

text = cell.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.7)
p_fmt.left_indent = Cm(3.8)
add = text.add_run("(полное наименование организации,\nосуществляющей образовательную деятельность)")
add.font.size = Pt(6)

text = cell.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(3.6)
add = text.add_run("Студенческий билет № ")
add.font.size = Pt(10)
add.bold = True
add = text.add_run(" 123456")
add.font.size = Pt(10)
add.bold = True
add.underline = True

text = cell.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(3.6)
add = text.add_run("Фамилия")
add.font.size = Pt(9)
add = text.add_run("Фамилия")
add.font.size = Pt(9)
add.underline = True


doc.save("resources/Студенческий.docx")
