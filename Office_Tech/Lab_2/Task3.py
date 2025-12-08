import docx
import docx.enum.text
from docx.shared import Cm, Mm, Pt
from docx.enum.text import WD_LINE_SPACING

personal = []

personal.append(input("Название программы ЭВМ: "))

personal.append(input("Фамилия Имя Отчество: "))

personal.append(input("Адрес прописки: "))

personal.append(input("Паспорт: "))

personal.append(input("Выдан: "))

personal.append(input("Кем: "))

personal.append(input("Фамилия И.О.: "))

doc = docx.Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

section = doc.sections[-1]
section.top_margin = Cm(2) #Верхний отступ
section.bottom_margin = Cm(2)#Нижний отступ
section.left_margin = Cm(3) #Отступ слева
section.right_margin = Cm(1.5) #Отступ справа

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(9.52)
add = text.add_run("В Федеральную службу\nпо интеллектуальной собственности\nБережковская наб., д. 30, корп. 1,\nг. Москва, Г-59, ГСП-3, 125993, \nРоссийская Федерация")
add.font.size = Pt(12)

void = doc.add_paragraph()
p_fmt = void.paragraph_format
p_fmt.first_line_indent = Cm(1)
add = void.add_run("")
add.font.size = Pt(13)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(6)
p_fmt.first_line_indent = Cm(1)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Название программы для ЭВМ или базы данных ")
add.font.size = Pt(12)
add = text.add_run(f"«{personal[0]}»")
add.font.size = Pt(12)
add.bold = True
add.underline = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(6)
p_fmt.first_line_indent = Cm(1)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("№  заявки ")
add.font.size = Pt(12)
add = text.add_run("__________________________________________________________")
add.font.size = Pt(12)
add.bold = True
add.underline = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(2)
p_fmt.left_indent = Cm(1.27)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run("(указывается при наличии регистрационного номера заявки)")
add.font.size = Pt(13)
add.font.superscript = True
add.font.italic = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(12)
p_fmt.left_indent = Cm(1)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run("Согласие на обработку персональных данных")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(12)
p_fmt.first_line_indent = Cm(1)
p_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Ф. И. О. субъекта персональных данных ")
add.font.size = Pt(12)
add = text.add_run(f"{personal[1]} _______________________________________________________________")
add.font.size = Pt(12)
add.bold = True
add.underline = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(6)
p_fmt.first_line_indent = Cm(1)
add = text.add_run("Адрес места жительства   ")
add.font.size = Pt(12)
add = text.add_run(f" {personal[2]}_____________")
add.font.size = Pt(12)
add.bold = True
add.underline = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(12)
p_fmt.first_line_indent = Cm(1.25)
p_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
add = text.add_run("Документ, удостоверяющий личность субъекта персональных данных, дата его \nвыдачи и выдавший орган ")
add.font.size = Pt(12)
add = text.add_run(f"_ паспорт {personal[3]}, выдан {personal[4]}\n{personal[5]}________________")
add.font.size = Pt(12)
add.bold = True
add.underline = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(12)
p_fmt.first_line_indent = Cm(1)
p_fmt.line_spacing = 1.2
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Подтверждаю согласие на обработку моих персональных данных, предусмотренную частью 3 статьи 3 Федерального закона от 27 июля 2006 г. № 152-ФЗ «О персональных данных», в целях предоставления Федеральной  службой  по интеллектуальной собственности  государственной услуги в соответствии с Федеральным законом от 27 июля 2010 г. № 210-ФЗ «Об организации предоставления государственных и муниципальных услуг».")
add.font.size = Pt(13)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1)
p_fmt.line_spacing = 1.2
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Мне известно, что в случае отзыва согласия на обработку персональных данных Федеральная служба по интеллектуальной собственности вправе продолжить обработку персональных данных без моего согласия в соответствии с частью 2 статьи 9, пунктом 4 части 1 статьи 6 Федерального закона от 27 июля 2006 г. № 152-ФЗ      «О персональных данных».")
add.font.size = Pt(13)

void = doc.add_paragraph()
p_fmt = void.paragraph_format
p_fmt.first_line_indent = Cm(1.27)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = void.add_run("")
add.font.size = Pt(12)

void = doc.add_paragraph()
p_fmt = void.paragraph_format
p_fmt.first_line_indent = Cm(1.27)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = void.add_run("")
add.font.size = Pt(12)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.4)
p_fmt.left_indent = Cm(0.5)
p_fmt.space_before = Pt(6)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Подпись					          ")
add.font.size = Pt(12)
add = text.add_run(f"/___{personal[6]}___________/")
add.font.size = Pt(12)
add.underline = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(10)
p_fmt.space_before = Pt(2)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("(Ф. И. О.  субъекта персональных данных)")
add.font.size = Pt(13)
add.font.subscript = True
add.font.italic = True

void = doc.add_paragraph()
p_fmt = void.paragraph_format
p_fmt.first_line_indent = Cm(10)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = void.add_run("")
add.font.size = Pt(12)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(9.75)
p_fmt.space_after = Pt(8)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Дата __________________________")
add.font.size = Pt(12)
add.underline = True

doc.save("resources/Согласие на обработку персональных данных.docx")