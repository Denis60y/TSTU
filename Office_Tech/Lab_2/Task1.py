import docx
import docx.enum.text
from docx.shared import Pt

doc = docx.Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'

head = doc.add_paragraph()
p_fmt = head.paragraph_format
p_fmt.space_after = Pt(10)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = head.add_run("Реферат")
add.bold = True
add.font.size = Pt(12)

void = doc.add_paragraph()
p_fmt = head.paragraph_format
p_fmt.space_after = Pt(10)

rows = ["Авторы", "Правообладатель", "Программа", "Аннотация", "Тип ЭВМ", "Язык", "ОС", "Объем программы"]

table = doc.add_table(rows=len(rows), cols=2)
table.style = 'Table Grid'

for i in range(len(rows)):
    if i == 0:
        col = ""
        while True:
            popa = input(f"{rows[i]} (Введите 'стоп' для остновки): ")
            if popa == "стоп":
                col = col[:-2]
                break
            else:
                col += popa + "," + "\n"
    else:
        col = input(f"{rows[i]}: ")

    table.cell(i, 0).text = rows[i]
    table.cell(i, 1).text = col
    
doc.save("resources/out/2.1/Реферат.docx")