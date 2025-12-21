import docx
import docx.enum.text
from docx.shared import Cm, Mm, Pt, Inches
from docx.enum.text import WD_LINE_SPACING
import csv
import os
from datetime import datetime


def save_to_csv(data, photos, text_photos, filename="lab_report_data.csv"):
    fields = [
        "Полное название организации",
        "Название кафедры",
        "№ лабораторной",
        "Название дисциплины",
        "Номер группы",
        "Фамилия И.О. студента",
        "Фамилия И.О. преподавателя",
        "Год",
        "Цель работы",
        "Вариант лабораторной",
        "Задание",
        "Решение",
        "Код программы",
        "Количество картинок",
        "Выводы",
        "Дата сохранения"
    ]
    
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    
    file_exists = os.path.isfile(filename)
    
    with open(filename, 'a', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
        
        if not file_exists:
            writer.writerow(fields)
        
        photos_info = ""
        for i in range(len(photos)):
            photos_info += f"Фото{i+1}:{photos[i]}; Подпись{i+1}:{text_photos[i]}"
            if i < len(photos) - 1:
                photos_info += " | "
        
        full_data = data + [len(photos)] + [photos_info] + [datetime.now().strftime("%Y-%m-%d %H:%M:%S")]
        
        writer.writerow(full_data)
    
    print(f"Данные лабораторной работы сохранены в файл: {filename}")


def is_valid_number(input_str, min_value=1, max_value=999):
    """Проверяет, является ли строка числом в заданном диапазоне"""
    if not input_str.isdigit():
        return False
    number = int(input_str)
    return min_value <= number <= max_value


personal = []
photos = []
text_photo = []

personal.append(input("Полное название организации: "))

personal.append(input("Название кафедры: "))

while True:
    lab_num = input("№ лабораторной: ")
    if is_valid_number(lab_num, 1, 100): 
        personal.append(lab_num)
        break
    else:
        print("Ошибка! Введите номер лабораторной (число от 1 до 100)")

personal.append(input("Название дисциплины: "))

personal.append(input("Номер группы: "))

personal.append(input("Фамилия И.О. студента: "))

personal.append(input("Фамилия И.О. преподавателя: "))

# Проверка года
current_year = datetime.now().year
while True:
    year_input = input("Год: ")
    if is_valid_number(year_input, 2000, current_year + 1):
        personal.append(year_input)
        break
    else:
        print(f"Ошибка! Введите корректный год (от 2000 до {current_year + 1})")

personal.append(input("Цель работы: "))

personal.append(input("Вариант лабораторной: "))

personal.append(input("Задание: "))

personal.append(input("Решение: "))

personal.append(input("Код программы: "))

while True:
    num_photos_input = input("Количество картинок: ")
    if is_valid_number(num_photos_input, 0, 20): 
        num_photos = int(num_photos_input)
        break
    else:
        print("Ошибка! Введите число от 0 до 20")

for i in range(num_photos):
    while True:
        photo_path = input(f"Путь к картинке номер {i+1}: ")
        if os.path.exists(photo_path):
            photos.append(photo_path)
            text_photo.append(f"Рисунок {i+1}. {input('Подпись к картинке: ')}")
            break
        else:
            print(f"Ошибка! Файл '{photo_path}' не найден. Проверьте путь.")

personal.append(input("Выводы: "))

save_to_csv(personal, photos, text_photo, "resources/out/2.1/lab_report_data.csv")

doc = docx.Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

section = doc.sections[-1]
section.top_margin = Cm(2) #Верхний отступ
section.bottom_margin = Cm(2)#Нижний отступ
section.left_margin = Cm(3) #Отступ слева
section.right_margin = Cm(1.5) #Отступ справа

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run(f"{personal[0]}")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("\n\n\n\n\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
add = text.add_run(f"{personal[1]}")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("\n\n\n\n\n\n\n\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run(f"Отчет по лабораторной работе №{personal[2]}")
add.font.size = Pt(24)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run(f"по дисциплине «{personal[3]}»")
add.font.size = Pt(16)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("\n")
add.font.size = Pt(16)
add = text.add_run("\n\n")
add.font.size = Pt(14)
add = text.add_run("\n\n\n\n")
add.font.size = Pt(18)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("                                                           ")
add.font.size = Pt(18)
add = text.add_run(f"Выполнил: студент группы {personal[4]}")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run(f"                                                                                                {personal[5]}")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run(f"                                                                            Проверил:  {personal[6]}")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("\n\n\n\n")
add.font.size = Pt(12)
add = text.add_run("\n\n\n\n\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run(f"Тамбов {personal[7]}")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.25)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Цель работы:")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.25)
add = text.add_run(f"{personal[8]}\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.25)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Задание:")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.25)
add = text.add_run(f"Вариант №{personal[9]}")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.25)
add = text.add_run(f"{personal[10]}\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("")
add.font.size = Pt(16)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.95)
add = text.add_run("Решение:")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.95)
add = text.add_run(f"{personal[11]}\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("Листинг программы")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("")
add.font.size = Pt(10)
add.font.name = "Courier New"

text = doc.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run(f"{personal[12]}")
add.font.size = Pt(12)
add.font.name = "Consolas"

text = doc.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("Результаты работы программы")
add.font.size = Pt(14)
add.bold = True

for i in range(len(photos)):
    try:
        para = doc.add_paragraph()
        p_fmt = para.paragraph_format
        p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run()
        run.add_picture(photos[i], width=Inches(5))
        
        caption = doc.add_paragraph()
        p_fmt = caption.paragraph_format
        p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        add = caption.add_run(f"{text_photo[i]}")
        add.font.size = Pt(12)
        
    except Exception as e:
        print(f"Ошибка при вставке изображения {photos[i]}: {e}")
        text = doc.add_paragraph()
        p_fmt = text.paragraph_format
        p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        add = text.add_run(f"Не удалось загрузить изображение: {photos[i]}")
        add.font.size = Pt(12)

        text = doc.add_paragraph()
        p_fmt = text.paragraph_format
        p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        add = text.add_run(f"{text_photo[i]}")
        add.font.size = Pt(12)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("")
add.font.size = Pt(14)
add.bold = True
add.font.name = "Courier New"

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.95)
add = text.add_run("")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.95)
add = text.add_run("Выводы:")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.95)
add = text.add_run(f"{personal[13]}")
add.font.size = Pt(14)

doc.save("resources/out/2.1/Структура отчета.docx")