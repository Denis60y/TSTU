import docx 
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

file_path = "C:/VScodeProjects/3-semestr/resources/input/Студак.xls"

def read_first_11_cells(file_path):
    try:
        df = pd.read_excel(file_path, header=None)
        
        first_row = df.iloc[0, :12] 
        
        result = first_row.tolist()
        
        while len(result) < 12:
            result.append(None)
            
        return result
        
    except FileNotFoundError:
        print(f"Ошибка: Файл {file_path} не найден")
        return None
    except Exception as e:
        print(f"Ошибка при чтении файла: {e}")
        return None

spisok = read_first_11_cells(file_path)
print(spisok)


def set_cell_borders(cell, top={"sz": 8, "val": "single", "color": "#000000"}, 
                     right={"sz": 8, "val": "single", "color": "#000000"},
                     bottom={"sz": 8, "val": "single", "color": "#000000"}, 
                     left={"sz": 8, "val": "single", "color": "#000000"}):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    
    for border, border_attrs in [('top', top), ('right', right), 
                                 ('bottom', bottom), ('left', left)]:
        border_element = OxmlElement(f'w:{border}')
        for key, value in border_attrs.items():
            border_element.set(qn(f'w:{key}'), str(value))
        tcBorders.append(border_element)
    
    tcPr.append(tcBorders)

doc = docx.Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'


table = doc.add_table(rows=1, cols=2)
table.style = 'TableGrid'
table.autofit = False  


row = table.rows[0]
row.height = Cm(6)


table.cell(0,0).width = Cm(3.6)  
table.cell(0,1).width = Cm(5.4)  

# Очищаем содержимое ячеек
for cell in row.cells:
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)


cell1 = row.cells[0]
set_cell_borders(cell1, top={"sz": 8, "val": "single", "color": "#000000"},
                 right={"sz": 0, "val": "none", "color": "#FFFFFF"},  # Прозрачная граница справа
                 bottom={"sz": 8, "val": "single", "color": "#000000"},
                 left={"sz": 8, "val": "single", "color": "#000000"})


cell2 = row.cells[1]
set_cell_borders(cell2, top={"sz": 8, "val": "single", "color": "#000000"},
                 right={"sz": 8, "val": "single", "color": "#000000"},
                 bottom={"sz": 8, "val": "single", "color": "#000000"},
                 left={"sz": 0, "val": "none", "color": "#FFFFFF"})  # Прозрачная граница слева

text = cell2.add_paragraph()
add = text.add_run("\nМинистерство образования и науки РФ")
add.font.size = Pt(5)
add.underline = True
add = text.add_run("___________________")
add.font.size = Pt(5)


text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.9)
add = text.add_run("(учредитель)")
add.font.size = Pt(6)

text = cell2.add_paragraph()
add = text.add_run("Федеральное государственное бюджетное образовательное учреждение")
add.underline = True
add.font.size = Pt(4)
add = text.add_run("____")
add.font.size = Pt(4)


text = cell2.add_paragraph()
add = text.add_run("Высшего образования «Тамбовский Государственный технический университет»")
add.font.size = Pt(4)
add.underline = True

text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.7)
p_fmt.left_indent = Cm(0.2)
add = text.add_run("(полное наименование организации,\nосуществляющей образовательную деятельность)")
add.font.size = Pt(5)

text = cell2.add_paragraph()
add = text.add_run("Студенческий билет № ")
add.font.size = Pt(10)
add.bold = True
add = text.add_run(f"{spisok[0]}_")
add.font.size = Pt(9)
add.bold = True
add.underline = True

text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(3)
add = text.add_run("Фамилия")
add.font.size = Pt(6)
add = text.add_run("_")
add.font.size = Pt(9)
add = text.add_run(f"{spisok[1]}")
add.font.size = Pt(9)
add.underline = True
add = text.add_run("_________________")
add.font.size = Pt(9)

text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(3)
add = text.add_run("Имя, отчество")
add.font.size = Pt(6)
add = text.add_run(f"_{spisok[2]}")
add.font.size = Pt(9)
add.underline = True
add = text.add_run("______")
add.font.size = Pt(9)


text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.9)
add = text.add_run("(последнее при наличии)")
add.font.size = Pt(6)

text = cell2.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("Форма обучения")
add.font.size = Pt(6)
add = text.add_run("_")
add.font.size = Pt(9)
add = text.add_run(f"{spisok[3]}")
add.font.size = Pt(9)
add.underline = True
add = text.add_run("_______________")
add.font.size = Pt(9)

text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(3)
add = text.add_run("Зачислен приказом от ")
add.font.size = Pt(6)
add = text.add_run(f"0{spisok[4]}.0{spisok[5]}")
add.font.size = Pt(7)
add.underline = True
add = text.add_run(" 20")
add.font.size = Pt(7)
add = text.add_run(f"{spisok[6]}")
add.font.size = Pt(7)
add.underline = True
add = text.add_run("г. №")
add.font.size = Pt(7)
add = text.add_run(f"{spisok[7]}")
add.font.size = Pt(7)
add.underline = True
add = text.add_run("_____")
add.font.size = Pt(7)


text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(3)
add = text.add_run("Дата выдачи ")
add.font.size = Pt(6)
add = text.add_run(f"« 0{spisok[8]} »")
add.font.size = Pt(8)
add.underline = True
add = text.add_run("______")
add.font.size = Pt(8)
add = text.add_run(f"0{spisok[9]}")
add.font.size = Pt(8)
add.underline = True
add = text.add_run("_______")
add.font.size = Pt(8)
add = text.add_run("20")
add.font.size = Pt(8)
add = text.add_run(f"{spisok[10]}")
add.font.size = Pt(8)
add.underline = True
add = text.add_run("г.")
add.font.size = Pt(8)

text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.space_before = Pt(3)
add = text.add_run("______________________")
add.font.size = Pt(7)

text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.3)
add = text.add_run("(подпись студента)")
add.font.size = Pt(4)


text = cell2.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run(f"_____{spisok[11]}_________________")
add.font.size = Pt(7)
add.underline = True

text = cell2.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(0.3)
add = text.add_run("(фамилия, имя, отчество(последнее - при наличии))")
add.font.size = Pt(4)

text = cell1.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("\n\n\n\n\n\n\n\n\n")
add.font.size = Pt(11)

text = cell1.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("М.П.")
add.font.size = Pt(6)

text = cell1.add_paragraph()
p_fmt = text.paragraph_format
add = text.add_run("Руководитель организацииn\nосуществляющей\nобразовательную\nдеятельность\nили иное уполномоченное\nим должностное лицо")
add.font.size = Pt(4)
add = text.add_run("                   ________")
add.font.size = Pt(6)

text = cell1.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(2.4)
add = text.add_run("(подпись)")
add.font.size = Pt(5)

doc.save("3-semestr/resources/out/2.2/Студенческий.docx")